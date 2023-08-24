import os
import zipfile
import uuid
from flask import Flask, request, send_file,jsonify, make_response
import PyPDF2
from docx2pdf import convert
from PyPDF2 import PdfMerger
from PyPDF2 import PdfReader, PdfWriter
# from python_docx2pdf import convert
from flask_restful import Api, Resource
from pymongo import MongoClient
from flask_cors import CORS
from pymongo import MongoClient
from pdf2docx import Converter as pdftodocxconverter
from docx import Document
import aspose.words as aw
from PyPDF2 import PdfReader, PdfWriter
from flask import Flask, request, render_template, send_file
from typing import Tuple
from pdf2docx import parse
import io
from io import BytesIO
from docx.shared import Inches
from flask import Flask, request,send_file, jsonify
from PyPDF2 import PdfMerger
from docx import Document
import pikepdf
from PIL import Image
import io
import fitz
import openpyxl
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
import tabula
import pandas as pd
import tempfile
# import requests
# import configparser

app = Flask(__name__)
CORS(app)
api = Api(app)

client = MongoClient(
    "mongodb+srv://darshan:darshan@cluster0.rdbg2qk.mongodb.net/?retryWrites=true&w=majority"
)  # Replace with your MongoDB connection URI
db = client["d2d"]  # Replace with your database name
collection = db["users"]  # Replace with your collection name

# @app.route("/register",method=['POST'])
# def new_user():
#     data=request.get_json()
#     name=data.get('name')
#     enrollment=data.get('enrollment')
#     student_data={
#         'name':name,
#         'enrollment':enrollment
#     }
#     collection.insert_one(student_data)
    # return ('data added')
    


@app.route("/convert", methods=["POST"])
def convert_file():
    file = request.files["file"]
    file.save("uploaded_file")
    try:
        if file.filename.lower().endswith(".pdf"):
            # cv=pdftodocxconverter('uploaded_file')
            # cv.convert('converted.docx')
            # cv.close
            pdftodocxconverter("uploaded_file").convert("converted.docx")
            pdftodocxconverter("uploaded_file").close
            return send_file("converted.docx", as_attachment=True)
        else:
            return "unsupported file format"
    except Exception as e:
        return str(e)


@app.route('/convertd2p', methods=['POST'])
def convert_word_to_pdf():
    word_file = request.files['file']
    
    # Save the uploaded Word file to a temporary location
    temp_file_path = 'temp_word_file.docx'
    word_file.save(temp_file_path)
    
    # Convert the Word file to PDF
    output_file_path = 'converted.pdf'
    convert(temp_file_path, output_file_path)
    
    # Clean up the temporary Word file
    os.remove(temp_file_path)
    
    # Send the converted PDF file as a response
    return send_file(
        output_file_path,
        as_attachment=True,
        download_name='converted.pdf'
    )
@app.route('/api/mergePDF', methods=['POST'])
def merge_pdf():
    file1 = request.files['file1']
    file2 = request.files['file2']

    pdf_merger = PdfMerger()
    pdf_merger.append(file1)
    pdf_merger.append(file2)

    output_buffer = io.BytesIO()
    pdf_merger.write(output_buffer)
    output_buffer.seek(0)

    return output_buffer.getvalue(), 200, {
        'Content-Type': 'application/pdf',
        'Content-Disposition': 'attachment; filename=merged.pdf',
    }





@app.route('/api/mergeDocx', methods=['POST'])
def merge_docx_endpoint():
    files = request.files.getlist('files')

    merged_filename = merge_docx(files)

    return send_file(merged_filename, as_attachment=True)

def merge_docx(files):
    merged_document = Document()

    for file in files:
        doc = Document(file)

        # Copy content
        for element in doc.element.body:
            merged_document.element.body.append(element)

        # Copy images
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                img = doc.part.rels[rel].target_part.blob
                img_stream = BytesIO(img)
                img_width, img_height = Inches(5.5), Inches(4)
                merged_document.add_picture(img_stream, width=img_width, height=img_height)

    merged_filename = 'merged.docx'
    merged_document.save(merged_filename)

    return merged_filename


@app.route('/compress', methods=['POST'])
def compress():
    file = request.files['file']
    temp_file_path = 'temp_word_file.pdf'
    file.save(temp_file_path)
    
    # Convert the Word file to PDF
    output_file_path = 'compressed1.pdf'
    compress_pdf(temp_file_path,output_file_path)
    import os
    os.remove(temp_file_path)
    return send_file(output_file_path, as_attachment=True,download_name='compressed.pdf')

def compress_pdf(input_file_path, output_file_path):
    pdf_reader = PdfReader(input_file_path)
    pdf_writer = PdfWriter()

    for page in pdf_reader.pages:
        page.compress_content_streams = True
        pdf_writer.add_page(page)

    with open(output_file_path, 'wb') as output_file:
        pdf_writer.write(output_file)


def split_pdf_at_page(input_file, split_page, output_file1, output_file2):
    input_pdf = io.BytesIO(input_file.read())  # Read the content of the uploaded file into a BytesIO object

    with open(output_file1, 'wb') as file1, open(output_file2, 'wb') as file2:
        reader = PyPDF2.PdfReader(input_pdf)
        total_pages = len(reader.pages)

        if split_page < 1 or split_page > total_pages:
            raise ValueError("Invalid page number. Please choose a page within the range 1 to {}".format(total_pages))

        # Create two PdfWriter objects to save the two split parts
        writer1 = PyPDF2.PdfWriter()
        writer2 = PyPDF2.PdfWriter()

        # Copy pages up to split_page into the first writer
        for page_num in range(split_page):
            page = reader.pages[page_num]
            writer1.add_page(page)

        # Copy pages from split_page onwards into the second writer
        for page_num in range(split_page, total_pages):
            page = reader.pages[page_num]
            writer2.add_page(page)

        # Write the split PDFs to output files
        writer1.write(file1)
        writer2.write(file2)

    # Create a zip archive containing the split PDFs
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        zip_file.write(output_file1)
        zip_file.write(output_file2)

    # Delete the split PDF files after creating the zip archive
    os.remove(output_file1)
    os.remove(output_file2)

    # Reset the buffer's file pointer to the beginning
    zip_buffer.seek(0)

    return zip_buffer

@app.route('/split', methods=['POST'])
def spliter():
    file = request.files['file']
    split_page = int(request.form["splitPage"])
    output_file1 = 'splittedp1.pdf'
    output_file2 = 'splittedp2.pdf'
    zip_buffer = split_pdf_at_page(file, split_page, output_file1, output_file2)

    # Set the filename for the downloaded zip file
    filename = 'split_pdf_files.zip'

    # Create the Flask response with the zip content
    response = make_response(zip_buffer.getvalue())
    response.headers['Content-Type'] = 'application/zip'
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'

    # Delete the split PDF files after creating the zip archive
    # os.remove(output_file1)
    # os.remove(output_file2)

    return response

@app.route('/excel2pdf', methods=['POST'])
def extopdf():
    file = request.files['file']
    output_file = 'translated.pdf'
    excel_to_pdf(file,output_file)
    return send_file(output_file, as_attachment=True,download_name='translated.pdf')


def excel_to_pdf(excel_file_path, pdf_file_path):
    try:
        # Load the Excel workbook
        workbook = openpyxl.load_workbook(excel_file_path)
        sheet = workbook.active

        # Get the data from the Excel sheet as a list of lists
        data = []
        for row in sheet.iter_rows():
            row_data = [cell.value for cell in row]
            data.append(row_data)

        # Create a PDF and add the data to it
        doc = SimpleDocTemplate(pdf_file_path, pagesize=landscape(letter))
        table = Table(data)

        # Set the style for the table
        style = TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)])
        table.setStyle(style)

        # Add the table to the PDF and build the document
        doc.build([table])

        return True
    except Exception as e:
        print(f"Error: {e}")
        return False
    
@app.route('/pdf2excel', methods=['POST'])
def pdftoex():
    file = request.files['file']
    excel_buffer = io.BytesIO()
    pdf_to_excel(file, excel_buffer)  # Pass the file object and Excel buffer
    excel_buffer.seek(0)
    return send_file(excel_buffer, as_attachment=True, download_name='converted.xlsx')

def pdf_to_excel(pdf_file, excel_buffer):
    try:
        # Create a temporary file to write the PDF content
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf_file:
            temp_pdf_file.write(pdf_file.read())
            temp_pdf_path = temp_pdf_file.name

        # Read the tables from the PDF file into a list of pandas DataFrames
        df_list = tabula.read_pdf(temp_pdf_path, pages='all', multiple_tables=True)

        # Create a Pandas Excel writer using XlsxWriter as the engine
        excel_writer = pd.ExcelWriter(excel_buffer, engine='xlsxwriter')

        # Write each DataFrame to a separate sheet in the Excel file
        for i, df in enumerate(df_list):
            df.to_excel(excel_writer, sheet_name=f'Sheet{i+1}', index=False)

        # Save the Excel file to the buffer
        excel_writer.save()

        # Delete the temporary PDF file
        os.remove(temp_pdf_path)

        return True
    except Exception as e:
        print(f"Error: {e}")
        return False


if __name__ == "__main__":
    app.run(debug=True,host="0.0.0.0",port=5000)



# nikeshapatel.it@svitvasad.ac.in